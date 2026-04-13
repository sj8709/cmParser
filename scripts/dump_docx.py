"""IBK input.docx 구조 덤프 — Step A 분석용."""
import sys
from pathlib import Path
from docx import Document
from docx.table import _Cell, Table

sys.stdout.reconfigure(encoding="utf-8")

DOCX = Path(__file__).parent.parent / "fixtures/ibk/input.docx"


def iter_cell_content(cell: _Cell, depth: int = 0):
    indent = "    " * depth
    for para in cell.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        bold = any(run.bold for run in para.runs if run.bold)
        marker = " [B]" if bold else ""
        print(f"{indent}P{marker}: {text[:120]}")
    for nt in cell.tables:
        print(f"{indent}[NESTED TABLE rows={len(nt.rows)} cols={len(nt.columns)}]")
        dump_table(nt, depth + 1, limit=3)


def dump_table(table: Table, depth: int = 0, limit: int | None = None):
    indent = "    " * depth
    rows = list(table.rows)
    for ri, row in enumerate(rows):
        if limit is not None and ri >= limit:
            print(f"{indent}... ({len(rows) - limit} more rows)")
            break
        print(f"{indent}Row {ri}:")
        for ci, cell in enumerate(row.cells):
            text_preview = cell.text.replace("\n", " | ")[:100]
            print(f"{indent}  C{ci}: {text_preview}")
            iter_cell_content(cell, depth + 2)


def main():
    doc = Document(str(DOCX))
    print(f"=== DOCX: {DOCX.name} ===")
    print(f"Total paragraphs (top-level): {len(doc.paragraphs)}")
    print(f"Total tables (top-level): {len(doc.tables)}")
    print()

    # 상위 단락 (제목 등)
    print("--- Top-level paragraphs (non-empty, first 30) ---")
    count = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        bold = any(run.bold for run in p.runs if run.bold)
        marker = " [B]" if bold else ""
        print(f"  P{marker}: {t[:150]}")
        count += 1
        if count >= 30:
            break
    print()

    # 각 테이블 요약 + 처음 3개 상세
    print("--- Tables summary ---")
    for ti, table in enumerate(doc.tables):
        print(f"Table[{ti}] rows={len(table.rows)} cols={len(table.columns)}")

    print()
    print("--- Tables detail (all, row preview) ---")
    for ti, table in enumerate(doc.tables):
        print(f"\n### Table[{ti}] rows={len(table.rows)} cols={len(table.columns)}")
        dump_table(table, depth=0, limit=8)


if __name__ == "__main__":
    main()
