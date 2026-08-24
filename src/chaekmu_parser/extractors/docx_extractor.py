"""
DOCX extractor — python-docx 기반.

규칙:
  - 모든 top-level 테이블 순서대로 추출
  - 각 셀: paragraphs(텍스트+bold 플래그) + 중첩 테이블 보존
  - Merged cell dedup: 같은 <w:tc> 엘리먼트가 여러 번 노출되는 경우 1회만 유지
  - text/is_bold는 단락 결합 결과로 채움 (하위 호환)
  - 단락의 Word 자동번호(w:numPr)는 텍스트에 번호가 없으므로 is_numbered 플래그로 보존
    (numFmt=bullet은 세부항목 글머리표이므로 제외)
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.table import _Cell as DocxCell

from chaekmu_parser.extractors.base import BaseExtractor
from chaekmu_parser.models import (
    RawCell,
    RawDocument,
    RawParagraph,
    RawRow,
    RawTable,
)


class DocxExtractor(BaseExtractor):

    @property
    def format_name(self) -> str:
        return "docx"

    def can_handle(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".docx"

    def extract(self, file_path: Path) -> RawDocument:
        doc = Document(str(file_path))
        self._num_formats = self._load_num_formats(doc)

        tables = [
            self._convert_table(t, source_index=idx)
            for idx, t in enumerate(doc.tables)
        ]

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        return RawDocument(
            source_path=str(file_path),
            format="docx",
            tables=tables,
            paragraphs=paragraphs,
        )

    def _convert_table(self, table: DocxTable, source_index: int) -> RawTable:
        rows: list[RawRow] = []
        for row in table.rows:
            seen_tc: set[int] = set()
            cells: list[RawCell] = []
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in seen_tc:
                    continue
                seen_tc.add(tc_id)
                cells.append(self._convert_cell(cell))
            rows.append(RawRow(cells=cells))
        return RawTable(rows=rows, source_index=source_index)

    def _convert_cell(self, cell: DocxCell) -> RawCell:
        paragraphs: list[RawParagraph] = []
        for p in cell.paragraphs:
            text = p.text
            is_bold = self._para_is_bold(p)
            is_numbered = self._para_is_numbered(p)
            paragraphs.append(
                RawParagraph(text=text, is_bold=is_bold, is_numbered=is_numbered)
            )

        joined_text = "\n".join(p.text for p in paragraphs)
        any_bold = any(p.is_bold and p.text.strip() for p in paragraphs)

        nested_tables: list[RawTable] = [
            self._convert_table(nt, source_index=-1)
            for nt in cell.tables
        ]

        return RawCell(
            text=joined_text,
            is_bold=any_bold,
            nested_tables=nested_tables,
            paragraphs=paragraphs,
        )

    @staticmethod
    def _load_num_formats(doc) -> dict[tuple[str, str], str]:
        """numbering.xml → {(numId, ilvl): numFmt}. 넘버링 파트가 없으면 빈 dict."""
        try:
            numbering = doc.part.numbering_part.element
        except (AttributeError, KeyError, NotImplementedError):
            return {}
        abstract_lvls: dict[str, dict[str, str]] = {}
        for a in numbering.findall(qn("w:abstractNum")):
            lvls: dict[str, str] = {}
            for lvl in a.findall(qn("w:lvl")):
                fmt = lvl.find(qn("w:numFmt"))
                if fmt is not None:
                    lvls[lvl.get(qn("w:ilvl"))] = fmt.get(qn("w:val"))
            abstract_lvls[a.get(qn("w:abstractNumId"))] = lvls
        result: dict[tuple[str, str], str] = {}
        for n in numbering.findall(qn("w:num")):
            aid = n.find(qn("w:abstractNumId"))
            if aid is None:
                continue
            for ilvl, fmt in abstract_lvls.get(aid.get(qn("w:val")), {}).items():
                result[(n.get(qn("w:numId")), ilvl)] = fmt
        return result

    def _para_is_numbered(self, p) -> bool:
        """단락 직접 지정 자동번호(numPr) 중 bullet이 아닌 것만 True."""
        pPr = p._p.pPr
        if pPr is None:
            return False
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            return False
        num_id = numPr.find(qn("w:numId"))
        if num_id is None or num_id.get(qn("w:val")) in (None, "0"):
            return False
        ilvl = numPr.find(qn("w:ilvl"))
        ilvl_val = ilvl.get(qn("w:val")) if ilvl is not None else "0"
        fmt = self._num_formats.get((num_id.get(qn("w:val")), ilvl_val))
        return fmt is not None and fmt != "bullet"

    @staticmethod
    def _para_is_bold(p, threshold: float = 0.5) -> bool:
        """단락을 'bold 제목'으로 볼지 글자 가중 다수결로 판정.

        `any(run.bold)`은 한두 글자(가운뎃점 등)만 우연히 bold여도 단락 전체를
        제목으로 승격시켜, 관리의무 세부항목이 제목으로 오추출되는 원인이 됐다.
        실제 글자 기준 과반이 bold일 때만 제목으로 본다 (잡티-bold/잡티-비bold 양방향 견고).
        """
        bold_chars = 0
        total_chars = 0
        for run in p.runs:
            stripped = run.text.strip()  # 공백-only run은 가중치에서 제외
            if not stripped:
                continue
            total_chars += len(stripped)
            if run.bold is True:
                bold_chars += len(stripped)
        return total_chars > 0 and (bold_chars / total_chars) >= threshold
