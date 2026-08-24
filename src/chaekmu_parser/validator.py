"""
5단계 정합성 검증 레이어 (HANDOFF §5.4).

1. **Stage 1 (재추출 비교)**: RawDocument가 원본을 충실히 담았는지 — 독립 경로로
   원본을 한 번 더 훑어 얻은 텍스트 세트와 raw 셀 텍스트를 비교.
2. **Stage 2 (parsed → raw substring)**: ParsedDocument의 각 verbatim 필드가
   raw 셀 텍스트 어딘가에 실제로 존재하는지 점검. 누락 = 후처리 왜곡 또는 파싱 오류.
3. **Stage 3 (역재조립 유사도)**: parsed 전체를 문자열로 재조립 후 raw와 유사도 측정.
   낮으면 대규모 누락/중복 가능성.
4. **Stage 4 (구조 정합성)**: bold-mode 문서에서 책무세부 ↔ 관리의무 제목의 1:1 대응을
   점검. 세부 0개·미매칭 제목 = 오추출 의심(warn), 대응 제목 없는 책무세부 = 원본 누락
   추정(info). 태그/번호 모드는 대응 관계가 없으므로 생략.
5. **Stage 5 (공통/고유 정합성)**: 책무표의 '임원 공통' 표기와 관리의무의 공통 책무
   분류가 임원별로 일치하는지 교차 검증. 분류 휴리스틱(마지막 블록 키워드)이
   놓친 공통 미분류/과분류를 원본 표기 기준으로 잡는다.

결과는 `ValidationReport`로 반환. GUI/테스트가 소비.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from difflib import SequenceMatcher
from pathlib import Path
from typing import Literal

from chaekmu_parser.models import (
    Executive,
    ParsedDocument,
    RawDocument,
    RawTable,
)
from chaekmu_parser.normalizer import _COMMON_RESP_MARKER

Severity = Literal["info", "warn", "error"]
Stage = Literal[1, 2, 3, 4, 5]

# ---------------------------------------------------------------------------
# 결과 자료구조
# ---------------------------------------------------------------------------
@dataclass(frozen=True)
class ValidationIssue:
    stage: Stage
    severity: Severity
    message: str
    context: str = ""  # 임원 직책, 필드명 등


@dataclass
class ValidationReport:
    issues: list[ValidationIssue] = field(default_factory=list)

    # Stage 1
    stage1_source_fragments: int = 0
    stage1_missing_fragments: int = 0

    # Stage 2
    stage2_verified_count: int = 0
    stage2_missing_count: int = 0

    # Stage 3
    stage3_similarity: float = 0.0

    # Stage 4 (구조 정합성)
    stage4_oversplit_count: int = 0
    stage4_missing_count: int = 0
    stage4_nonidentical_count: int = 0

    # Stage 5 (공통/고유 정합성)
    stage5_common_mismatch_count: int = 0

    @property
    def passed(self) -> bool:
        return not any(i.severity == "error" for i in self.issues)

    @property
    def has_warnings(self) -> bool:
        return any(i.severity == "warn" for i in self.issues)

    def counts_by_severity(self) -> dict[Severity, int]:
        result: dict[Severity, int] = {"info": 0, "warn": 0, "error": 0}
        for i in self.issues:
            result[i.severity] += 1
        return result

    def summary_line(self) -> str:
        c = self.counts_by_severity()
        status = "✓ 통과" if self.passed else "❌ 실패"
        return (
            f"{status} — 오류 {c['error']} · 경고 {c['warn']} · 정보 {c['info']} | "
            f"Stage1 누락 {self.stage1_missing_fragments}/"
            f"{self.stage1_source_fragments} · "
            f"Stage2 확인 {self.stage2_verified_count} (누락 {self.stage2_missing_count}) · "
            f"Stage3 유사도 {self.stage3_similarity:.1%} · "
            f"Stage4 오추출의심 {self.stage4_oversplit_count}·원본누락 {self.stage4_missing_count}"
            f"·비동일 {self.stage4_nonidentical_count} · "
            f"Stage5 공통불일치 {self.stage5_common_mismatch_count}"
        )

    def summary_block(self) -> str:
        """로그/뷰어용 여러 줄 요약 — Stage별로 개행해 가독성 확보."""
        c = self.counts_by_severity()
        status = "✓ 통과" if self.passed else "❌ 실패"
        return "\n".join([
            f"🔍 검증 {status} — 오류 {c['error']} · 경고 {c['warn']} · 정보 {c['info']}",
            f"   ├ Stage 1 재추출 비교  : 원본 {self.stage1_source_fragments}단락 중 누락 {self.stage1_missing_fragments}",
            f"   ├ Stage 2 raw 대조     : 확인 {self.stage2_verified_count} · 누락 {self.stage2_missing_count}",
            f"   ├ Stage 3 재조립 유사도 : {self.stage3_similarity:.1%}",
            f"   ├ Stage 4 구조 정합성   : 오추출의심 {self.stage4_oversplit_count} · 원본누락 {self.stage4_missing_count} · 비동일 {self.stage4_nonidentical_count}",
            f"   └ Stage 5 공통/고유     : 불일치 {self.stage5_common_mismatch_count}",
        ])


# ---------------------------------------------------------------------------
# 공개 API
# ---------------------------------------------------------------------------
# Stage별 임계치 — 본 임계치 이하로 떨어지면 경고/오류
_STAGE1_MISSING_RATIO_WARN = 0.01   # 1% 이상 누락
_STAGE1_MISSING_RATIO_ERROR = 0.05  # 5% 이상 누락 → 오류
_STAGE3_SIMILARITY_WARN = 0.55
_STAGE3_SIMILARITY_ERROR = 0.30
_MAX_MISSING_ISSUE_PER_STAGE = 10
# Stage 4 원본 누락 추정: 임원 한 명에서 이 건수 이상 뭉텅이로 빠지면 원본이 안 적은
# 게 아니라 파싱 누락(예: 자동번호 제목이 세부항목으로 흡수)일 가능성이 높아 warn 승격.
_STAGE4_MISSING_PER_EXEC_WARN = 3


def validate(
    parsed: ParsedDocument, raw: RawDocument, source_path: Path | None = None
) -> ValidationReport:
    """5단계 검증 실행."""
    report = ValidationReport()
    _run_stage1(parsed, raw, source_path, report)
    _run_stage2(parsed, raw, report)
    _run_stage3(parsed, raw, report)
    _run_stage4(parsed, report)
    _run_stage5(parsed, report)
    return report


# ---------------------------------------------------------------------------
# Stage 1
# ---------------------------------------------------------------------------
def _run_stage1(
    parsed: ParsedDocument,
    raw: RawDocument,
    source_path: Path | None,
    report: ValidationReport,
) -> None:
    source_fragments = _extract_source_fragments(source_path, raw.format)
    if source_fragments is None:
        # 지원하지 않는 포맷은 Stage 1 생략 (예: Phase 2 HWP 전)
        return
    raw_joined = _raw_text_blob(raw)
    missing: list[str] = [f for f in source_fragments if f not in raw_joined]
    report.stage1_source_fragments = len(source_fragments)
    report.stage1_missing_fragments = len(missing)

    if not source_fragments:
        return
    ratio = len(missing) / len(source_fragments)
    if ratio >= _STAGE1_MISSING_RATIO_ERROR:
        report.issues.append(ValidationIssue(
            1, "error",
            f"Stage 1 — 원본의 {ratio:.1%} 단락이 raw에 없음. Extractor 오류 가능성 높음"
        ))
    elif ratio >= _STAGE1_MISSING_RATIO_WARN:
        report.issues.append(ValidationIssue(
            1, "warn",
            f"Stage 1 — 원본의 {ratio:.1%} 단락이 raw에 없음"
        ))
    for m in missing[:_MAX_MISSING_ISSUE_PER_STAGE]:
        report.issues.append(ValidationIssue(
            1, "warn", f"Stage 1 — 누락 단락: {_truncate(m, 80)}"
        ))


def _extract_source_fragments(
    source_path: Path | None, fmt: str
) -> list[str] | None:
    """원본을 독립 경로로 훑어 non-empty stripped 단락 리스트 반환. 지원 안하면 None."""
    if source_path is None or not source_path.exists():
        return None
    if fmt == "docx":
        return _docx_fragments(source_path)
    # hwp/pdf는 Phase 2/3에서 확장
    return None


def _docx_fragments(path: Path) -> list[str]:
    """python-docx로 모든 단락·테이블·중첩테이블을 한 번 더 훑는다 (DocxExtractor와 독립)."""
    from docx import Document
    from docx.table import Table as DocxTable

    doc = Document(str(path))
    fragments: list[str] = []

    def walk_cell(cell) -> None:
        for p in cell.paragraphs:
            txt = p.text.strip()
            if txt:
                fragments.append(txt)
        for t in cell.tables:
            walk_table(t)

    def walk_table(t: DocxTable) -> None:
        seen: set[int] = set()
        for row in t.rows:
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in seen:
                    continue
                seen.add(tc_id)
                walk_cell(cell)

    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            fragments.append(t)
    for table in doc.tables:
        walk_table(table)
    return fragments


def _raw_text_blob(raw: RawDocument) -> str:
    """RawDocument의 모든 셀 텍스트와 top-level 단락을 하나의 문자열로."""
    parts: list[str] = list(raw.paragraphs)
    for t in raw.tables:
        parts.extend(_collect_cell_text(t))
    return "\n".join(parts)


def _collect_cell_text(t: RawTable) -> list[str]:
    out: list[str] = []
    for row in t.rows:
        for cell in row.cells:
            if cell.text:
                out.append(cell.text)
            for nested in cell.nested_tables:
                out.extend(_collect_cell_text(nested))
    return out


# ---------------------------------------------------------------------------
# Stage 2
# ---------------------------------------------------------------------------
def _run_stage2(
    parsed: ParsedDocument, raw: RawDocument, report: ValidationReport
) -> None:
    """각 임원의 verbatim 필드가 raw에 실제로 존재하는지 확인."""
    raw_blob = _raw_text_blob(raw)
    verified = 0
    missing_details: list[tuple[str, str, str]] = []  # (exec, field, value)

    for e in parsed.executives:
        for field_name, value in _verbatim_fields(e):
            if not value:
                continue
            if value in raw_blob:
                verified += 1
            else:
                missing_details.append((e.position.replace("\n", ", "), field_name, value))

    report.stage2_verified_count = verified
    report.stage2_missing_count = len(missing_details)

    if not missing_details:
        return

    report.issues.append(ValidationIssue(
        2, "warn",
        f"Stage 2 — raw에서 찾을 수 없는 parsed 값 {len(missing_details)}건"
    ))
    for pos, fname, val in missing_details[:_MAX_MISSING_ISSUE_PER_STAGE]:
        report.issues.append(ValidationIssue(
            2, "warn",
            f"Stage 2 — {fname} = {_truncate(val, 60)}",
            context=pos,
        ))


def _verbatim_fields(e: Executive) -> list[tuple[str, str]]:
    """원본과 글자 단위로 동일해야 하는 필드 목록.

    라벨 매칭 위해 trimming은 했지만 법령 '등' 제거 같은 후처리된 필드는 제외.
    """
    fields: list[tuple[str, str]] = [
        ("name", e.name),
        ("title", e.title),
        ("appointed_date", e.appointed_date),
        ("concurrent_yn", e.concurrent_yn),
        ("departments", e.departments),
    ]
    for idx, c in enumerate(e.committees):
        fields.append((f"committee[{idx}].name", c.name))
        fields.append((f"committee[{idx}].role", c.role))
    for idx, r in enumerate(e.responsibilities):
        fields.append((f"responsibility[{idx}].category", r.category))
        if r.details:
            fields.append((f"responsibility[{idx}].details", r.details[0]))
    for idx, o in enumerate(e.obligations):
        fields.append((f"obligation[{idx}].category", o.category))
        for i, item in enumerate(o.items):
            fields.append((f"obligation[{idx}].items[{i}]", item))
    return fields


# ---------------------------------------------------------------------------
# Stage 3
# ---------------------------------------------------------------------------
def _run_stage3(
    parsed: ParsedDocument, raw: RawDocument, report: ValidationReport
) -> None:
    raw_blob = _raw_text_blob(raw)
    parsed_blob = _parsed_text_blob(parsed)
    # quick_ratio는 set-based 빠른 근사. 정확한 ratio는 O(n²)이라 대용량 부적합.
    matcher = SequenceMatcher(None, raw_blob, parsed_blob, autojunk=True)
    sim = matcher.quick_ratio()
    report.stage3_similarity = sim

    if sim < _STAGE3_SIMILARITY_ERROR:
        report.issues.append(ValidationIssue(
            3, "error",
            f"Stage 3 — 재조립 유사도 {sim:.1%} — 대규모 누락/중복 의심"
        ))
    elif sim < _STAGE3_SIMILARITY_WARN:
        report.issues.append(ValidationIssue(
            3, "warn",
            f"Stage 3 — 재조립 유사도 {sim:.1%} — 검토 권장"
        ))


def _parsed_text_blob(parsed: ParsedDocument) -> str:
    parts: list[str] = []
    for e in parsed.executives:
        parts.extend([e.position, e.name, e.title, e.appointed_date])
        parts.append(e.concurrent_detail)
        parts.append(e.departments)
        for c in e.committees:
            parts.extend([c.name, c.role, c.cycle, c.matters])
        parts.append(e.responsibility_summary)
        parts.append(e.assign_date)
        for r in e.responsibilities:
            parts.append(r.category)
            parts.extend(r.details)
            parts.append(r.raw_law_reg)
        for o in e.obligations:
            parts.extend([o.category, *o.items])
    return "\n".join(p for p in parts if p)


# ---------------------------------------------------------------------------
# Stage 4 — 구조 정합성 (책무세부 ↔ 관리의무 제목 1:1 대응)
# ---------------------------------------------------------------------------
# 매칭 임계치: 공통책무의 어미 차이("관리책임" vs "관리에 대한 책임", ≈0.96)는
# 통과시키고, 진짜 누락(대응 없음, 실측 IBK ≈0.43)만 잡도록 KDB/IBK 실측으로 0.7 선정.
_STAGE4_MATCH_RATIO = 0.7


def _run_stage4(parsed: ParsedDocument, report: ValidationReport) -> None:
    """bold-mode 문서에서 책무세부와 관리의무 제목의 1:1 대응을 점검.

    - **오추출 의심(warn)**: 세부항목 0개 + 어떤 책무세부와도 미매칭인 관리의무 제목.
      세부 한 줄이 bold 잡티로 제목으로 오승격된 경우(예: KDB IT부문장).
    - **원본 누락 추정(info)**: 대응하는 관리의무 제목이 없는 책무세부.
      원본이 해당 책무의 관리의무를 적지 않은 경우(예: IBK 동산/부동산).
      단, 한 임원에서 `_STAGE4_MISSING_PER_EXEC_WARN`건 이상이면 파싱 누락 의심으로
      집계 이슈를 **warn**으로 승격(예: 라이나 2026-07 Word 자동번호 제목 흡수).
    - **비동일 대응(info)**: 대응되긴 하나 책무세부와 관리의무 책무명이 글자 단위로
      다른 경우. ICR은 둘을 정확 일치 키로 조인하므로 비동일이면 다운스트림 조인이
      깨진다. 어미 차이(예: '관리책임' vs '관리에 대한 책임')부터 원문 자체의 표현
      차이(예: '퇴직연금 상품' vs '퇴직상품')까지 사람이 검토하도록 전부 나열.

    gating: 관리의무에 세부항목이 하나도 없으면(태그/번호 모드 — 제목이 책무세부와
    대응하지 않음) 검사를 생략한다. fail-soft: 모두 warn/info라 저장은 막지 않는다.
    """
    oversplit: list[ValidationIssue] = []
    missing: list[ValidationIssue] = []
    nonidentical: list[ValidationIssue] = []
    missing_heavy: list[str] = []  # 임원당 누락 건수가 임계치 이상인 직책

    for e in parsed.executives:
        resp_details = [d for r in e.responsibilities for d in r.details]
        if not resp_details or not e.obligations:
            continue
        if not any(o.items for o in e.obligations):
            continue  # bold-mode 아님 → 제목-책무세부 대응 검사 부적합
        obl_titles = [o.category for o in e.obligations]
        pos = e.position.replace("\n", ", ")

        for o in e.obligations:
            if o.items:
                continue  # 세부항목 있으면 정상 제목
            if not _stage4_has_match(o.category, resp_details):
                oversplit.append(ValidationIssue(
                    4, "warn",
                    f"Stage 4 — 관리의무 제목 오추출 의심: "
                    f"{_truncate(o.category, 50)} (세부항목 0개·책무세부 미매칭)",
                    context=pos,
                ))

        missing_before = len(missing)
        for d in resp_details:
            if not _stage4_has_match(d, obl_titles):
                missing.append(ValidationIssue(
                    4, "info",
                    f"Stage 4 — 원본 관리의무 누락 추정: 책무세부 "
                    f"{_truncate(d, 50)} 에 대응하는 관리의무 없음",
                    context=pos,
                ))
                continue
            # 대응은 되지만 글자 단위로 동일하지 않은 쌍 → 정확 일치 조인 깨짐 후보
            best = _stage4_best_match(d, obl_titles)
            if best is not None and best != d:
                raw_ratio = SequenceMatcher(None, d, best).ratio()
                nonidentical.append(ValidationIssue(
                    4, "info",
                    f"Stage 4 — 책무세부↔관리의무 책무명 비동일(유사도 {raw_ratio:.0%}): "
                    f"책무세부 «{_truncate(d, 70)}» ↔ 관리의무 «{_truncate(best, 70)}»",
                    context=pos,
                ))
        exec_missing = len(missing) - missing_before
        if exec_missing >= _STAGE4_MISSING_PER_EXEC_WARN:
            missing_heavy.append(f"{pos} {exec_missing}건")

    report.stage4_oversplit_count = len(oversplit)
    report.stage4_missing_count = len(missing)
    report.stage4_nonidentical_count = len(nonidentical)

    if oversplit:
        report.issues.append(ValidationIssue(
            4, "warn", f"Stage 4 — 관리의무 제목 오추출 의심 {len(oversplit)}건"
        ))
        report.issues.extend(oversplit[:_MAX_MISSING_ISSUE_PER_STAGE])
    if missing_heavy:
        report.issues.append(ValidationIssue(
            4, "warn",
            f"Stage 4 — 원본 관리의무 누락 추정 {len(missing)}건 — 임원당 "
            f"{_STAGE4_MISSING_PER_EXEC_WARN}건 이상 뭉텅이 누락 {len(missing_heavy)}명 "
            f"(파싱 누락 의심: {', '.join(missing_heavy[:5])})"
        ))
    elif missing:
        report.issues.append(ValidationIssue(
            4, "info", f"Stage 4 — 원본 관리의무 누락 추정 {len(missing)}건"
        ))
        report.issues.extend(missing[:_MAX_MISSING_ISSUE_PER_STAGE])
    if nonidentical:
        report.issues.append(ValidationIssue(
            4, "info",
            f"Stage 4 — 책무세부↔관리의무 책무명 비동일 {len(nonidentical)}건 "
            f"(ICR 정확 일치 조인 대상 — 검토 필요)"
        ))
        report.issues.extend(nonidentical[:_MAX_MISSING_ISSUE_PER_STAGE])


def _stage4_norm(s: str) -> str:
    """가운뎃점/불릿 통일 + 공백 제거 — 어미·표기 차이를 줄여 매칭 안정화."""
    s = s.replace("∙", "·").replace("ㆍ", "·").replace("•", "·")
    return re.sub(r"\s+", "", s)


def _stage4_has_match(target: str, candidates: list[str]) -> bool:
    """target이 candidates 중 하나와 동일/포함/유사(ratio≥임계치)면 True."""
    nt = _stage4_norm(target)
    if not nt:
        return True  # 빈 값은 검사 제외
    for c in candidates:
        nc = _stage4_norm(c)
        if not nc:
            continue
        if nt == nc or nt in nc or nc in nt:
            return True
        if SequenceMatcher(None, nt, nc).ratio() >= _STAGE4_MATCH_RATIO:
            return True
    return False


def _stage4_best_match(target: str, candidates: list[str]) -> str | None:
    """target과 정규화 유사도가 가장 높은 candidate 반환 (동일/포함은 최우선)."""
    nt = _stage4_norm(target)
    if not nt:
        return None
    best: str | None = None
    best_ratio = -1.0
    for c in candidates:
        nc = _stage4_norm(c)
        if not nc:
            continue
        ratio = 1.0 if (nt == nc or nt in nc or nc in nt) else SequenceMatcher(None, nt, nc).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best = c
    return best


# ---------------------------------------------------------------------------
# Stage 5 — 공통/고유 정합성 (책무 '임원 공통' 표기 ↔ 공통 관리의무)
# ---------------------------------------------------------------------------
# 책무(RESP) 표의 '임원 공통' 마커는 normalizer가 단일 출처로 정의 — 분류와
# 검증이 같은 정규식을 공유해야 휴리스틱·교차검증의 의미가 일치한다.


def _run_stage5(parsed: ParsedDocument, report: ValidationReport) -> None:
    """책무표의 '임원 공통' 표기와 관리의무의 공통 책무 분류가 일치하는지 교차 검증.

    분류 로직(normalizer)은 마지막 블록 키워드 휴리스틱에 의존해 공통/고유를 정한다.
    이 단계는 원본이 명시한 '임원 공통' 표기를 독립 근거로 삼아 휴리스틱의 오분류를 잡는다.

    - **공통 미분류 의심(warn)**: 책무에 '임원 공통' 표기가 있으나 관리의무는 전부 고유.
    - **공통 과분류 의심(info)**: 관리의무에 공통 책무가 있으나 책무엔 '임원 공통' 표기 없음.

    대표이사처럼 둘 다 없는 경우는 정상으로 통과한다.
    """
    mismatches: list[ValidationIssue] = []
    for e in parsed.executives:
        resp_common = any(
            _COMMON_RESP_MARKER.search(r.category or "") for r in e.responsibilities
        )
        oblig_common = any(o.type == "공통 책무" for o in e.obligations)
        if resp_common == oblig_common:
            continue
        pos = e.position.replace("\n", ", ")
        if resp_common and not oblig_common:
            mismatches.append(ValidationIssue(
                5, "warn",
                "Stage 5 — 책무에 '임원 공통' 표기가 있으나 관리의무가 전부 고유로 분류됨 "
                "(공통 미분류 의심)",
                context=pos,
            ))
        else:  # oblig_common and not resp_common
            mismatches.append(ValidationIssue(
                5, "info",
                "Stage 5 — 관리의무에 공통 책무가 있으나 책무에 '임원 공통' 표기 없음 "
                "(분류 근거 확인 권장)",
                context=pos,
            ))

    report.stage5_common_mismatch_count = len(mismatches)
    report.issues.extend(mismatches[:_MAX_MISSING_ISSUE_PER_STAGE])


# ---------------------------------------------------------------------------
# 유틸
# ---------------------------------------------------------------------------
def _truncate(text: str, limit: int) -> str:
    if len(text) <= limit:
        return text
    return text[:limit] + "…"
